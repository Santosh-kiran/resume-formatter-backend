from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def proper_case(name):
    parts = name.strip().split()
    if len(parts) >= 2:
        return parts[0].capitalize() + " " + parts[1].capitalize()
    return name.capitalize()

def clean_text(text):
    text = re.sub(r'[•●▪►]', '', text)
    text = re.sub(r'http\S+', '', text)
    return text.strip()

def generate_docx(text):
    doc = Document()

    lines = [l.strip() for l in text.split("\n") if l.strip()]

    candidate_name = proper_case(lines[0])
    filename = f"{candidate_name}.docx"
    output_path = f"outputs/{filename}"

    # Candidate Name
    name_para = doc.add_paragraph()
    run = name_para.add_run(candidate_name)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    sections = {
        "Summary": [],
        "Technical Skills": [],
        "Education, Certification & Training": [],
        "Professional Experience": []
    }

    current_section = None

    for line in lines:
        if line in sections:
            current_section = line
            continue
        if current_section:
            sections[current_section].append(line)

    for section, content in sections.items():
        doc.add_paragraph("")
        heading = doc.add_paragraph()
        h_run = heading.add_run(section)
        h_run.bold = True
        h_run.font.name = "Times New Roman"
        h_run.font.size = Pt(10)

        for line in content:
            cleaned = clean_text(line)
            if cleaned:
                para = doc.add_paragraph(f"  • {cleaned}")
                para.runs[0].font.name = "Times New Roman"
                para.runs[0].font.size = Pt(10)

    doc.save(output_path)
    return output_path, filename
