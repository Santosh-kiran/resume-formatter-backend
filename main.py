from fastapi import FastAPI, UploadFile, File
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
import shutil
import uuid
import os

app = FastAPI()

# Enable CORS for frontend (Vercel)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # You can restrict later
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
def root():
    return {"status": "THIS IS NEW VERSION"}


@app.post("/upload")
async def upload_resume(file: UploadFile = File(...)):

    if not file.filename.endswith(".docx"):
        return JSONResponse(
            status_code=400,
            content={"error": "Only .docx files are supported"}
        )

    # Save uploaded file
    input_path = f"/tmp/{uuid.uuid4()}.docx"

    with open(input_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    try:
        original_doc = Document(input_path)
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": "Invalid DOCX file"}
        )

    # Extract text
    content = [p.text.strip() for p in original_doc.paragraphs if p.text.strip()]

    if not content:
        return JSONResponse(
            status_code=400,
            content={"error": "Resume appears empty"}
        )

    # Create new formatted document
    new_doc = Document()

    # Add Name
    new_doc.add_heading(content[0], level=0)

    # Simple section detection
    summary = []
    skills = []
    experience = []
    education = []

    for line in content[1:]:
        lower = line.lower()

        if "skill" in lower:
            skills.append(line)
        elif "experience" in lower:
            experience.append(line)
        elif "education" in lower:
            education.append(line)
        else:
            summary.append(line)

    # Add Sections Properly
    if summary:
        new_doc.add_heading("Professional Summary", level=1)
        for s in summary[:5]:
            new_doc.add_paragraph(s)

    if skills:
        new_doc.add_heading("Technical Skills", level=1)
        for s in skills:
            new_doc.add_paragraph(s)

    if experience:
        new_doc.add_heading("Professional Experience", level=1)
        for e in experience:
            new_doc.add_paragraph(e)

    if education:
        new_doc.add_heading("Education", level=1)
        for ed in education:
            new_doc.add_paragraph(ed)

    # Save formatted resume
    output_path = f"/tmp/formatted_{uuid.uuid4()}.docx"
    new_doc.save(output_path)

    return FileResponse(
        output_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="Formatted_Resume.docx"
    )

